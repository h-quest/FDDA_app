# -*- coding: utf-8 -*-
# 3S Solar Plus - Master Project 
# Author: Hugo Quest (hugo.quest@hotmail.com)

import numpy as np
import pandas as pd
import datetime as dt
import math
from suntime import Sun
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from sklearn.metrics import auc
import ipywidgets as widgets
from IPython.display import clear_output
import pvlib
import streamlit as st
import rdtools
import statsmodels.api as sm
import os
from matplotlib import gridspec
from matplotlib.dates import DateFormatter
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots

import warnings
warnings.simplefilter('ignore')
from FDD_preprocessing import *

######################

def Get_module_data():

    ### Import module data from available database

    cec_modules = pvlib.pvsystem.retrieve_sam(name='CECMod')
    cec_module = cec_modules.Canadian_Solar_Inc__CS5P_220M

    ### Create copy and change parameters to match 3S module

    threeS_module = cec_module.copy()
    threeS_module['Technology'] = 'PERC'
    threeS_module['STC'] = 195
    threeS_module['A_c'] = 1300*pow(10,-3)*875*pow(10,-3)
    threeS_module['Length'] = 1300*pow(10,-3)
    threeS_module['Width'] = 875*pow(10,-3)
    threeS_module['N_s'] = 40
    threeS_module['I_sc_ref'] = 9.5
    threeS_module['V_oc_ref'] = 26.9
    threeS_module['I_mp_ref'] = 9.1
    threeS_module['V_mp_ref'] = 21.9
    threeS_module['alpha_sc'] = 0.0405
    threeS_module['beta_oc'] = -0.2943
    #threeS_module['T_NOCT'] = 
    #threeS_module['a_ref'] = 
    threeS_module['I_L_ref'] = 9.1
    #threeS_module['I_o_ref'] =  
    threeS_module['R_s'] = 1.1969
    threeS_module['R_sh_ref'] = 5002.3687
    #threeS_module['Adjust'] = 
    threeS_module['gamma_r'] = -0.375
    threeS_module['BIPV'] = 'Yes'
    threeS_module['Version'] = 'V1 Hugo Quest - MSc Thesis'
    threeS_module['Date'] = '04/2021'
    threeS_module['Name'] = 'Megaslate II L'
    threeS_module.name = 'Megaslate II L'

    d = {k: threeS_module[k] for k in ['a_ref', 'I_L_ref', 'I_o_ref', 'R_sh_ref', 'R_s']}
    
    return threeS_module, d

######################

def Clear_sky_single_diode(data, module, d, meta, tilt_=None, azimuth_=None, freq='10min'):
    
    latitude=meta['Latitude'].values[0]

    longitude=meta['Longitude'].values[0]
    
    if(tilt_==None):
        tilt=meta['Tilt1'].values[0]
        if(tilt==np.nan):
            tilt=meta['Tilt2'].values[0]
        if(tilt==np.nan):
            tilt=meta['Tilt3'].values[0]
    else:
        tilt=tilt_
        
    if(azimuth_==None):
        azimuth=meta['Orient1'].values[0]
        if(azimuth==np.nan):
            azimuth=meta['Orient2'].values[0]
        if(azimuth==np.nan):
            azimuth=meta['Orient3'].values[0]
    else:
        azimuth=azimuth_

    temp_model='open_rack_glass_polymer'
    
    df_cs = []
    df_cs = data.copy()
    #df_cs.set_index('Time', inplace=True)
    df_cs = df_cs[(df_cs.T != float(0)).any()]

    # TZ is required for irradiance transposition
    #df_cs.index = df_cs.index.tz_localize('UTC', ambiguous = 'infer')
    df_cs = df_cs.reset_index()

    # Resampling
    df_cs = df_cs.resample(freq, on='Time').mean()

    # Calculate POA irradiance from DHI, GHI inputs
    loc = pvlib.location.Location(latitude,longitude, tz = 'UTC')
    sun = loc.get_solarposition(df_cs.index)

    # Calculate the clear sky POA irradiance
    clearsky = loc.get_clearsky(df_cs.index, solar_position = sun, model="simplified_solis")
    cs_sky = pvlib.irradiance.isotropic(tilt, clearsky.dhi)
    cs_beam = pvlib.irradiance.beam_component(tilt, azimuth, sun.zenith, sun.azimuth, clearsky.dni)
    df_cs['clearsky_poa'] = cs_beam + cs_sky

    # Compute cell temperature
    from pvlib.temperature import TEMPERATURE_MODEL_PARAMETERS

    params = TEMPERATURE_MODEL_PARAMETERS['sapm'][temp_model]

    # Calculate the clearsky temperature
    df_cs['clearsky_Tamb'] = rdtools.get_clearsky_tamb(df_cs.index, latitude, longitude)
    df_clearsky_temp = pvlib.temperature.sapm_cell(df_cs.clearsky_poa, df_cs.clearsky_Tamb, 0, **params)
    df_cs['clearsky_Tcell'] = df_clearsky_temp

    # Compute single diode model parameters
    photocurrent, saturation_current, resistance_series, resistance_shunt, nNsVth = (
        pvlib.pvsystem.calcparams_desoto(df_cs.clearsky_poa,
                                         df_cs.clearsky_Tcell,
                                         module['alpha_sc'],
                                         EgRef=1.121,
                                         dEgdT=-0.0002677, **d))

    # Simulate single diode model outputs
    single_diode_out = pvlib.pvsystem.singlediode(photocurrent, saturation_current,
                                                  resistance_series, resistance_shunt, nNsVth)

    # Scale the simulation to the studied system
    from sklearn.preprocessing import MinMaxScaler
    scaler_1 = MinMaxScaler(feature_range=(0,1))
    scaler_2 = MinMaxScaler(feature_range=(0,1))
    scaler_3 = MinMaxScaler(feature_range=(0,1))
    scaler_4 = MinMaxScaler(feature_range=(0,1))
    scaler_5 = MinMaxScaler(feature_range=(0,1))
    scaler_6 = MinMaxScaler(feature_range=(0,1))
    df_cs['Power_sc'] = scaler_1.fit_transform(np.array(df_cs['Power']).reshape(-1, 1))
    df_cs['Voltage_sc'] = scaler_3.fit_transform(np.array(df_cs['Voltage']).reshape(-1, 1))
    df_cs['Current_sc'] = scaler_5.fit_transform(np.array(df_cs['Current']).reshape(-1, 1))
    single_diode_out['p_mp_sc'] = scaler_2.fit_transform(np.array(single_diode_out['p_mp']).reshape(-1, 1))
    single_diode_out['v_mp_sc'] = scaler_4.fit_transform(np.array(single_diode_out['v_mp']).reshape(-1, 1))
    single_diode_out['i_mp_sc'] = scaler_6.fit_transform(np.array(single_diode_out['i_mp']).reshape(-1, 1))
    single_diode_out['p_mp'] = scaler_1.inverse_transform(np.array(single_diode_out['p_mp_sc']).reshape(-1, 1))
    single_diode_out['v_mp'] = scaler_3.inverse_transform(np.array(single_diode_out['v_mp_sc']).reshape(-1, 1))
    single_diode_out['i_mp'] = scaler_5.inverse_transform(np.array(single_diode_out['i_mp_sc']).reshape(-1, 1))

    # Align the two signals
    from scipy import signal

    def align(day):
        plot_data_1 = df_cs[df_cs.index.date==day].sort_index().dropna()
        plot_data_2 = single_diode_out[single_diode_out.index.date==day].sort_index().dropna()
        s0 = plot_data_2['v_mp'].values
        s1 = plot_data_1['Voltage'].values

        if((s0.size != 0) and (s1.size != 0)):
            dx = np.mean(np.diff(plot_data_2.index))
            shift = np.argmax(signal.correlate(s0, s1, mode = 'valid', method = 'direct'))*dx
            x = plot_data_2.index.values
            x1 = plot_data_1.index.values
            Initial_shift = x1[0]- x[0]
            final_shift = Initial_shift - shift
            res = final_shift
        else: res = np.nan
        return res

    days = []
    shifts = []
    for day_ in np.unique(df_cs.index.date):
        days.append(day_)
        shifts.append(align(day_))

    shift_pd = pd.DataFrame({'Day': days, 'Shift':shifts})
    df_cs['Date'] = df_cs.index.date
    df_cs = df_cs.reset_index()
    df_cs = df_cs.merge(shift_pd , left_on = 'Date', right_on = 'Day')
    df_cs['Shifted_time'] = df_cs['Time']-df_cs['Shift']
    df_cs = df_cs.set_index('Shifted_time')
    
    return df_cs, single_diode_out

######################

def Plot_clear_sky(df_cs, single_diode_out):

    plot_data_1 = df_cs.sort_index()
    plot_data_2 = single_diode_out.sort_index()

    s0 = plot_data_2['p_mp'].values
    s1 = plot_data_1['Power'].values

    x = plot_data_2.index
    x1 = plot_data_1.index

    fig = go.Figure([go.Scatter(x=x, y=s0, line_color='darkorange', name='Clear-sky')])
    fig.add_trace(go.Scatter(mode="lines", x=x1, y=s1, name='Actual', line_color='steelblue'))

    fig.update_xaxes(
        rangeslider_visible=True,
        rangeselector=dict(
            buttons=list([
                dict(count=6, label="6m", step="month", stepmode="backward"),
                dict(count=1, label="1m", step="month", stepmode="backward"),
                dict(count=7, label="7d", step="day", stepmode="backward"),
                dict(step="all")
            ])
        )
    )
    
    fig.update_layout(
        paper_bgcolor='rgba(0, 0, 0, 0)')
    
    fig.update_yaxes(title_text='Power [W]', title_font_size=12)
    
    #fig.show(config={'toImageButtonOptions': {'format': 'svg', 'filename': 'saved_image','width': 800,'scale': 1}})
    
    return fig
    
######################

def Get_merged_df(single_diode_out_cs, df_cs, meta, method='clear-sky-kt'):
    
    if((method=='clear-sky') | (method=='clear-sky-kt')):
        single_diode_out = single_diode_out_cs
        df = df_cs
    elif(method=='solcast'):
        single_diode_out = single_diode_out_sc
        df = df_sc
    
    df_merged = df.merge(single_diode_out, left_on=df.index, right_on=single_diode_out.index)
    df_merged.drop('key_0', axis=1, inplace=True)
    df_merged.set_index('Time', inplace=True)
    df_merged['Time'] = df_merged.index

    from suntime import Sun

    latitude=meta['Latitude'].values[0] 
    longitude=meta['Longitude'].values[0]

    sun = Sun(latitude, longitude)

    def get_sunrise(row, sun):
        return (sun.get_local_sunrise_time(row.name))
    def get_sunset(row, sun):
        return (sun.get_local_sunset_time(row.name))

    merged['sunrise'] = merged.apply(lambda x: get_sunrise(x,sun), axis=1)
    merged['sunset'] = merged.apply(lambda x: get_sunset(x,sun), axis=1)
    merged['hour'] = merged.reset_index().apply(lambda x: x['Time'].time(), axis = 1).values
    
    v_mp_std = merged.copy()
    v_mp_std = v_mp_std.groupby(v_mp_std.index.date).agg(np.std, ddof=0)['v_mp']
    v_mp_std.name = 'v_mp_std'
    merged['Time'] = merged.index
    merged = merged.merge(v_mp_std, left_on=merged.index.date, right_on=v_mp_std.index)
    merged.set_index('Time', inplace=True)
    merged.drop(['key_0'], axis=1, inplace=True)
    
    v_std = merged.copy()
    v_std = v_std.groupby(v_std.index.date).agg(np.std, ddof=0)['Voltage']
    v_std.name = 'v_std'
    merged['Time'] = merged.index
    merged = merged.merge(v_std, left_on=merged.index.date, right_on=v_std.index)
    merged.set_index('Time', inplace=True)
    merged.drop(['key_0'], axis=1, inplace=True)
    
    i_mp_std = merged.copy()
    i_mp_std = i_mp_std.groupby(i_mp_std.index.date).agg(np.std, ddof=0)['i_mp']
    i_mp_std.name = 'i_mp_std'
    merged['Time'] = merged.index
    merged = merged.merge(i_mp_std, left_on=merged.index.date, right_on=i_mp_std.index)
    merged.set_index('Time', inplace=True)
    merged.drop(['key_0'], axis=1, inplace=True)
    
    i_std = merged.copy()
    i_std = i_std.groupby(i_std.index.date).agg(np.std, ddof=0)['Current']
    i_std.name = 'i_std'
    merged['Time'] = merged.index
    merged = merged.merge(i_std, left_on=merged.index.date, right_on=i_std.index)
    merged.set_index('Time', inplace=True)
    merged.drop(['key_0'], axis=1, inplace=True)
    
    if(method=='clear-sky-kt'):
        
        merged['kt'] = merged['Power']/merged['p_mp']
        merged['kt'][merged['kt']>1] = 1

        test = merged.copy()
        test = test.dropna()
        test = test.groupby(test.index.date).agg(list)[['kt']]
        test['mean'] = test['kt'].apply(lambda x: np.mean(np.array(x)))
        test['var'] = test['kt'].apply(lambda x: np.mean(np.abs(np.diff(x))))

        def Weather_class(row):
            mean = row['mean']
            var = row['var']
            if(0.6-mean>var):
                return 0 # overcast
            elif(-0.6+0.9*mean>=var):
                return 2 # cloudless
            else:
                return 1 # broken clouds

        test['weather'] = test.apply(Weather_class, axis=1)

        merged['Time'] = merged.index
        merged = merged.merge(test, left_on=merged.index.date, right_on=test.index).drop('key_0', axis=1)
        merged = merged.set_index('Time')
        merged = merged[~merged.index.duplicated(keep='first')]
    
    if(method=='clear-sky'):
        
        poa_sc = df_sc['poa_global']
        merged = merged.merge(poa_sc, left_on=merged.index, right_on=poa_sc.index).rename(columns={'key_0':'Time'}).set_index('Time')
        merged['kt'] = merged['poa_global']/merged['clearsky_poa']
        merged['ktp'] = merged['Power']/merged['p_mp']
        merged['kt'][merged['kt']>1] = 1
        merged['ktp'][merged['ktp']>1] = 1

        test = merged.copy()
        test = test.dropna()
        test = test.groupby(test.index.date).agg(list)[['kt', 'ktp']]
        test['mean'] = test['kt'].apply(lambda x: np.mean(np.array(x)))
        test['var'] = test['kt'].apply(lambda x: np.mean(np.abs(np.diff(x))))
        test['mean_p'] = test['ktp'].apply(lambda x: np.mean(np.array(x)))
        test['var_p'] = test['ktp'].apply(lambda x: np.mean(np.abs(np.diff(x))))

        def Weather_class(row):
            mean = row['mean']
            var = row['var']
            if(0.6-mean>var):
                return 0 # overcast
            elif(-0.72+0.8*mean>=var):
                return 2 # cloudless
            else:
                return 1 # broken clouds

        test['weather'] = test.apply(Weather_class, axis=1)

        merged['Time'] = merged.index
        merged = merged.merge(test, left_on=merged.index.date, right_on=test.index).drop('key_0', axis=1)
        merged = merged.set_index('Time')
        merged = merged[~merged.index.duplicated(keep='first')]
        
    return merged

######################

def Diagnose_Fault(row, low_factor=1, high_factor=1, v_std=None):
    
    voltage = row['Voltage']
    if(v_std==None):
        v_std = row['v_std']
    else:
        v_std = v_std
    v_mp = row['v_mp'] + high_factor*row['v_std']
    v_mp_low = row['v_mp'] - low_factor*v_std
    current = row['Current']
    i_mp = row['i_mp']
    i_mp_low = row['i_mp'] - 0.5*row['i_std']
    sunrise = (row['sunrise']+dt.timedelta(minutes=20)).time()
    sunset = (row['sunset']-dt.timedelta(minutes=30)).time()
    time = row.name.time()
    
    if(np.isin(row.index, 'weather').any()):
        
        weather = row['weather']
        
        if(sunrise < time < sunset):

            if(v_mp_low <= voltage <= v_mp): ## --- Normal
                return 0
            elif((voltage < v_mp_low) & ((weather==0)|(weather==1))): 
                if(current < i_mp_low):
                    return 2 ## --- Low voltage and low current
                else:
                    return 3 ## -- Low voltage
            elif(((voltage < v_mp_low) | (current < i_mp_low)) & ((weather==1))): ## --- Cloudy
                return 1
            elif((voltage > v_mp) & ((weather==0)|(weather==1))): 
                if(current < i_mp_low):
                    return 4 ## --- High voltage and low current
                else:
                    return 5 ## --- High voltage
            elif(((voltage > v_mp) | (current < i_mp_low)) & ((weather==2))): ## --- Cloudy 
                return 1
        else:
            if(voltage>0):
                return np.nan
        
    else:     
        
        cloudy = row['cloudy']

        if(sunrise < time < sunset):

            if(v_mp_low <= voltage <= v_mp): ## --- Normal
                return 0
            elif((voltage < v_mp_low) & (cloudy==0)): 
                if(current < i_mp_low):
                    return 2 ## --- Low voltage and low current
                else:
                    return 3 ## -- Low voltage
            elif(((voltage < v_mp_low) | (current < i_mp_low)) & (cloudy==1)): ## --- Cloudy
                return 1
            elif((voltage > v_mp) & (cloudy==0)): ## --- Shading 2
                if(current < i_mp_low):
                    return 4 ## --- High voltage and low current
                else:
                    return 5 ## --- High voltage
            elif(((voltage > v_mp) | (current < i_mp_low)) & (cloudy==1)): ## --- Cloudy 
                return 1
        else:
            if(voltage>0):
                return np.nan
        
######################

def Plot_fault_heatmap(merged):
    
    df_heatmap = merged.groupby(merged.index.date).agg(list)
    df_heatmap = df_heatmap[1:-1]

    y = df_heatmap.index
    z = df_heatmap['Fault'].values
    df_heatmap['length_hour'] = df_heatmap['hour'].apply(lambda x: len(x))
    maxval = df_heatmap.groupby('length_hour').count()
    maxval = maxval[maxval['Power']==maxval['Power'].max()].index[0]
    dates = df_heatmap['hour'][df_heatmap['length_hour']==maxval][0]

    fig = go.Figure(data=go.Heatmap(
            z=z,
            x=dates,
            y=y,
            #colorscale='Viridis',
            colorscale=[[0,'rgb(68,1,84)'],[0.15,'rgb(68,1,84)'],
              [0.15,'white'], [0.17,'white'],
              [0.17,'grey'],[0.32,'grey'],
              [0.32,'white'], [0.34,'white'],   
              [0.34,'steelblue'],[0.49,'steelblue'],
              [0.49,'white'], [0.51,'white'],
              [0.51,'rgb(57,86,143)'],[0.66,'rgb(57,86,143)'],
              [0.66,'white'], [0.68,'white'],
              [0.68,'rgb(31,150,139)'], [0.83,'rgb(31,150,139)'],
              [0.83,'white'], [0.85,'white'],
              [0.85,'teal'], [1,'teal']], 
            colorbar=dict(
                title="<b>Fault Type</b>",
                titleside="top",
                #tickvals=[0.1, 1, 2, 2.9],
                tickvals=[0.4, 1.2, 2.1, 2.9, 3.8, 4.6],
                ticktext=['F0 - Normal', 'F1 - Cloudy', 'F2 - Low V, low I', 'F3 - Low V', 'F4 - High V, low I', 
                          'F5 - High V'], ticks='')
                #ticks="outside"),
                #colorscale=[[0, "green"], [0.25, "green"], [0.25, "orange"], [0.5, "orange"], 
                #            [0.5, 'tomato'], [0.75, 'tomato'], [0.75, 'darkgrey'],  [1, "darkgrey"]]
                #colorscale=[(0, "midnightblue"), (0.33, "darkorange"), (0.66, "tomato"), (1, "darkgrey")]
    ))
    

    fig.update_layout(
        #title='Shading Faults Heatmap',
        xaxis_nticks=6,
        yaxis_nticks=6,
        width=895, 
        height=530,
        paper_bgcolor='rgba(0, 0, 0, 0)',
        font=dict(size=15),
        template='ggplot2')
    
    fig.update_xaxes(tickmode = 'array',
                     tickvals = [dt.time(4,0),dt.time(8,0),dt.time(12,0),dt.time(16,0),dt.time(20,0)],
                     mirror=True)
    fig.update_yaxes(mirror=True)

    fig.show(config={'toImageButtonOptions': {'format': 'svg', 'filename': 'saved_image','width': 750,'scale': 1}})
    
    return fig


######################

def Plot_voltage_compare(merged, low_factor=1, high_factor=1, v_std=None, loss_variable='Current'):

    from plotly.subplots import make_subplots

    fig = make_subplots(rows=3, cols=1, row_width=[0.1, 0.1, 0.8], shared_xaxes=True, 
                        specs=[[{"secondary_y": True}],[{"secondary_y": False}],[{"secondary_y": False}]],
                       vertical_spacing=0.02)

    plot_data = merged.copy()
    
    last_date = plot_data.iloc[len(merged)-1].name.date()
    plot_data = plot_data[plot_data.index.date<last_date]

    s0 = plot_data['Power'].values
    s1 = plot_data['Voltage'].values
    s2 = plot_data['v_mp'].values + high_factor*plot_data['v_std']
    s3 = plot_data['p_mp'].values
    s5 = plot_data['Current'].values
    s6 = plot_data['i_mp'].values
    s7 = plot_data['i_mp'].values - 0.5*plot_data['i_std'].values
    s2[s3<10] = 0
    s7[s6<0.1] = 0
    s7[s7<0] = 0
    if(v_std==None):
        s4 = plot_data['v_mp'].values - low_factor*plot_data['v_std'].values
    else:
        s4 = plot_data['v_mp'].values - low_factor*v_std
    x = plot_data.index

    s4[s4<0] = 0
    s2[s4==0] = 0
    s4[s2==0] = 0
    
    if(loss_variable=='Current'):
        fig.add_trace(go.Scatter(mode="lines", x=x, y=s5, name='I-dc', line_color='steelblue'), row=1, col=1)
        fig.add_trace(go.Scatter(mode="lines", x=x, y=s6, name='I-sd', line_color='orange', line=dict(width=0.3), showlegend=False), row=1, col=1)
        fig.add_trace(go.Scatter(mode="lines", x=x, y=s7, name='I-sd', line_color='orange',line=dict(width=0.3), fill='tonexty'), row=1, col=1)
    elif(loss_variable=='Power'):
        fig.add_trace(go.Scatter(mode="lines", x=x, y=s0, name='P-dc', line_color='steelblue'), row=1, col=1)
        fig.add_trace(go.Scatter(mode="lines", x=x, y=s3, name='P-sd', line_color='orange'), row=1, col=1)
    fig.add_trace(go.Scatter(mode="lines", x=x, y=s1, name='V-dc', line_color='darkslategrey'),secondary_y=True, row=1, col=1)
    fig.add_trace(go.Scatter(mode="lines", x=x, y=s2, name='V-sd', line_color='grey', line=dict(width=0.3), showlegend=False),secondary_y=True, row=1, col=1)
    fig.add_trace(go.Scatter(mode="lines", x=x, y=s4, name='V-sd', line_color='grey',line=dict(width=0.3), fill='tonexty'),secondary_y=True, row=1, col=1)

    y = ['Fault Type ']
    dates = merged.index
    list_ = []
    list_.append(np.nan_to_num(merged['Fault'].values, copy=True, nan=0.0, posinf=None, neginf=None))
    z = np.array(list_)
    
    
    fig.add_trace(go.Heatmap(z=z,x=dates,y=y,
              colorscale=[[0,'rgb(68,1,84)'],[0.15,'rgb(68,1,84)'],
              [0.15,'white'], [0.17,'white'],
              [0.17,'grey'],[0.32,'grey'],
              [0.32,'white'], [0.34,'white'],   
              [0.34,'steelblue'],[0.49,'steelblue'],
              [0.49,'white'], [0.51,'white'],
              [0.51,'rgb(57,86,143)'],[0.66,'rgb(57,86,143)'],
              [0.66,'white'], [0.68,'white'],
              [0.68,'rgb(31,150,139)'], [0.83,'rgb(31,150,139)'],
              [0.83,'white'], [0.85,'white'],
              [0.85,'teal'], [1,'teal']],                                  
                colorbar=dict(
                title="<b>Fault Type</b>",
                xanchor='left',
                yanchor='bottom',
                len=0.5, 
                titleside="top",
                tickvals=[0.4, 1.2, 2.1, 2.9, 3.8, 4.6],
                ticktext=['F0 - Normal', 'F1 - Cloudy', 'F2 - Low V, low I', 'F3 - Low V', 'F4 - High V, low I', 
                          'F5 - High V'], ticks='')), row=2, col=1)

    dates = merged.index
    list1_ = []
    
    if(loss_variable=='Current'):
        y1 = ['Current loss']
        list1_.append(np.nan_to_num(merged['I_diff'].values, copy=True, nan=0.0, posinf=None, neginf=None))
        title_ = "<b>Current loss</b>"
        title_yaxis = r'$\text{Current [A]}$'
    elif(loss_variable=='Power'):
        y1 = ['Power loss ']
        list1_.append(np.nan_to_num(merged['P_diff'].values, copy=True, nan=0.0, posinf=None, neginf=None))
        title_ = "<b>Power loss</b>"
        title_yaxis = r'$\text{Power [W]}$'
    
    z1 = np.array(list1_)
    
    fig.add_trace(go.Heatmap(z=z1,x=dates,y=y1,colorscale='RdYlGn_r',
                    colorbar=dict(
                    title=title_,
                    xanchor='left',
                    yanchor='top',
                    len=0.5, 
                    titleside='top',
                    ticks=''))
                    , row=3, col=1)

    fig.update_layout(legend=dict(
        orientation="h",
        yanchor="bottom",
        y=1.02,
        xanchor="right",
        x=1
    ))
    
    fig.update_xaxes(
        mirror=True,
        rangeslider_visible=True,
        rangeselector=dict(
            buttons=list([
                dict(count=6, label="6m", step="month", stepmode="backward"),
                dict(count=1, label="1m", step="month", stepmode="backward"),
                dict(count=7, label="7d", step="day", stepmode="backward"),
                dict(step="all")
            ])
        )
    )
    fig.update_yaxes(title_text=title_yaxis, title_font_size=12, secondary_y=False, row=1, col=1)
    fig.update_yaxes(title_text=r'$\text{Voltage [V]}$', title_font_size=12, secondary_y=True,showgrid=False, row=1, col=1)
    fig.update_yaxes(mirror=True)
    fig.update_layout(template='ggplot2', paper_bgcolor='rgba(0, 0, 0, 0)')
    
    config = {
      'toImageButtonOptions': {
        'format': 'svg', # one of png, svg, jpeg, webp
        'filename': 'custom_image',
        'height': 500,
        'width': 900,
        'scale': 1
      }
    }
    
    fig.show(config=config)
    return fig

######################

def Power_diff(row):
    
    if((row['Fault']!=0)&(not np.isnan(row['Fault']))):
        power=row['Power']
        p_mp=row['p_mp']
        diff = (p_mp-power)/p_mp
        if(diff<0):
            return 0
        else:
            return diff
    elif(np.isnan(row['Fault'])):
        return np.nan
    else:
        return 0
    
######################

def Power_diff_full(row):
    
    if(np.isnan(row['Fault'])):
        return np.nan
    else:
        power=row['Power']
        p_mp=row['p_mp']
        diff = (p_mp-power)/p_mp
        if(diff<0):
            return 0
        else:
            return diff
        
######################

def Current_diff_full(row):
    
    if(np.isnan(row['Fault'])):
        return np.nan
    else:
        current=row['Current']
        i_mp=row['i_mp']
        diff = (i_mp-current)/i_mp
        if(diff<0):
            return 0
        else:
            return diff

######################        
        
def Plot_Pdiff_heatmap(merged):

    df_heatmap = merged.groupby(merged.index.date).agg(list)
    df_heatmap = df_heatmap[1:-1]

    y = df_heatmap.index
    z = df_heatmap['P_diff'].values
    df_heatmap['length_hour'] = df_heatmap['hour'].apply(lambda x: len(x))
    maxval = df_heatmap.groupby('length_hour').count()
    maxval = maxval[maxval['Power']==maxval['Power'].max()].index[0]
    dates = df_heatmap['hour'][df_heatmap['length_hour']==maxval][0]

    fig = go.Figure(data=go.Heatmap(
            z=z,
            x=dates,
            y=y,
            colorscale='RdYlGn_r',
            colorbar=dict(
                title="<b>Power loss</b>",
                titleside="top",
                ticks="outside")
    ))

    fig.update_layout(
        #title='Shading Faults Heatmap',
        xaxis_nticks=6,
        yaxis_nticks=6,
        width=895, 
        height=530,
        paper_bgcolor='rgba(0, 0, 0, 0)',
        font=dict(size=15),
        template='ggplot2')
    
    fig.update_xaxes(tickmode = 'array',
                 tickvals = [dt.time(4,0),dt.time(8,0),dt.time(12,0),dt.time(16,0),dt.time(20,0)])

    fig.show(config={'toImageButtonOptions': {'format': 'svg', 'filename': 'saved_image','width': 750,'scale': 1}})
    
    return fig

######################    

def Plot_Idiff_heatmap(merged):

    df_heatmap = merged.groupby(merged.index.date).agg(list)
    df_heatmap = df_heatmap[1:-1]

    y = df_heatmap.index
    z = df_heatmap['I_diff'].values
    df_heatmap['length_hour'] = df_heatmap['hour'].apply(lambda x: len(x))
    maxval = df_heatmap.groupby('length_hour').count()
    maxval = maxval[maxval['Power']==maxval['Power'].max()].index[0]
    dates = df_heatmap['hour'][df_heatmap['length_hour']==maxval][0]
    

    fig = go.Figure(data=go.Heatmap(
            z=z,
            x=dates,
            y=y,
            colorscale='RdYlGn_r',
            colorbar=dict(
                title="<b>Current loss</b>",
                titleside="top",
                ticks="outside")
    ))

    fig.update_layout(
        #title='Shading Faults Heatmap',
        xaxis_nticks=6,
        yaxis_nticks=6,
        width=895, 
        height=530,
        paper_bgcolor='rgba(0, 0, 0, 0)',
        font=dict(size=15),
        template='ggplot2')
    
    fig.update_xaxes(tickmode = 'array',
                 tickvals = [dt.time(4,0),dt.time(8,0),dt.time(12,0),dt.time(16,0),dt.time(20,0)])

    fig.show(config={'toImageButtonOptions': {'format': 'svg', 'filename': 'saved_image','width': 750,'scale': 1}})
    
    return fig

######################    

def relative_time(row):
    sunrise = row['sunrise'].hour*60 + row['sunrise'].minute
    sunset = row['sunset'].hour*60 + row['sunset'].minute
    time = row['hour'].hour*60 + row['hour'].minute
    relative_time = (time-sunrise)/(sunset-sunrise)
    return relative_time

######################

def Plot_fault_distribution(merged):

    merged['relative_time'] = merged.apply(relative_time, axis=1)
    faulty_hours = merged[(merged['Fault'].isin([2,3,4,5]))]
    faulty_hours = faulty_hours.sort_values(by='relative_time')
    f2 = faulty_hours[faulty_hours['Fault']==2]['relative_time'].values
    f3 = faulty_hours[faulty_hours['Fault']==3]['relative_time'].values
    f4 = faulty_hours[faulty_hours['Fault']==4]['relative_time'].values
    f5 = faulty_hours[faulty_hours['Fault']==5]['relative_time'].values

    fig = go.Figure(data=[go.Histogram(x=f2, nbinsx=100, bingroup=1, 
                                       name='F2 - Low V, low I', marker_color='rgb(57,86,143)', opacity=0.8,
                                       xbins=dict(start=0, end=1, size=0.01))])
    fig.add_trace(go.Histogram(x=f3, nbinsx=100, bingroup=1, 
                                       name='F3 - Low V', marker_color='steelblue', opacity=0.8,
                                       xbins=dict(start=0, end=1, size=0.01)))
    fig.add_trace(go.Histogram(x=f4, nbinsx=100, bingroup=1, 
                                       name='F4 - High V, low I', marker_color='rgb(31,150,139)', opacity=0.8,
                                       xbins=dict(start=0, end=1, size=0.01)))
    fig.add_trace(go.Histogram(x=f5, nbinsx=100, bingroup=1, 
                                       name='F5 - High V', marker_color='teal', opacity=0.8,
                                       xbins=dict(start=0, end=1, size=0.01)))

    fig.update_xaxes(tickmode = 'array',
                     tickvals = [0.2, 0.5, 0.8],
                     ticktext= ['0.2 (morning)', ' 0.5 (noon)', '0.8 (afternoon)'])

    fig.update_layout(
        xaxis_range=[0, 1],
        barmode='stack',
        xaxis_title_text=r'$\text{Relative time of day}$', # xaxis label
        yaxis_title_text=r'$\text{Fault count}$', # yaxis label
        bargap=0, # gap between bars of adjacent location coordinates
        template='simple_white',
        width=900,
        paper_bgcolor='rgba(0, 0, 0, 0)')

    fig.show(config={'toImageButtonOptions': {'format': 'svg', 'filename': 'saved_image','width': 700,'scale': 1}})
    
    return fig  

######################

def FFDA_run_all(data_file, metadata_file, year, freq='30min'):

    datetime_format = '%d/%m/%Y %H:%M'
    data = pd.read_csv(data_file)
    data.rename(columns={'Unnamed: 0': 'Time'}, inplace=True)
    data['Time'] = pd.to_datetime(data['Time'], format=datetime_format)
    data['Power'] = data['Power'].astype(float)
    data['Voltage'] = data['Voltage'].astype(float)
    data['Current'][data['Voltage']==0] = 0
    data['Current'] = data['Current'].astype(float)
    
    df = data.copy()
    meta = pd.read_csv(metadata_file)

    
    ### Select studied year
    df = df[df['Time'].dt.year==year]
    
    ### PV module data
    module, d = Get_module_data()
    
    
    ### Filtering data
    df = Filter_data(df, meta)
    
    ### Clear-sky single-diode model 
    df_cs, single_diode_out_cs = Clear_sky_single_diode(df, module, d, meta, tilt_=meta['Tilt'].values[0], 
                                                 azimuth_=meta['Orientation'].values[0], freq=freq)
    
 
    ### Plot Clear-sky simulation
    clear_sky_fig = Plot_clear_sky(df_cs, single_diode_out_cs)
    
    ### Merge all data for analysis
    merged = Get_merged_df(single_diode_out_cs, df_cs, meta, method='clear-sky-kt')
    
    ### Fault diagnosis 
    merged['Fault'] = merged.apply(lambda x: Diagnose_Fault(x, low_factor=2, high_factor=0, v_std=merged['v_std'].max()), axis=1)
    merged['P_diff'] = merged.apply(Power_diff_full, axis=1)
    merged['I_diff'] = merged.apply(Current_diff_full, axis=1)
    
  
    ### Fault Heatmap figure 
    fault_heatmap_fig = Plot_fault_heatmap(merged)
    
   
    ### Fault level heatmap (Current)
    current_heatmap_fig = Plot_Idiff_heatmap(merged)

    ### Fault level heatmap (Power)
    power_heatmap_fig = Plot_Pdiff_heatmap(merged)
    
    
    ### Plot voltage and power to compare with simulations
    voltage_compare_fig = Plot_voltage_compare(merged, low_factor=2, high_factor=0, v_std=merged['v_std'].max(), loss_variable='Current')
    
    ### Shading Factor and Percentage
    merged['Shading_Loss'] = np.where((merged['Fault']!=0)&(merged['Fault']!=5)&(merged['Fault']==merged['Fault']), merged['P_diff'], 0)
    merged['Shading_Loss'][merged['Fault']!=merged['Fault']] = np.nan
    Fault_losses = (merged['Shading_Loss'].dropna()[merged['Shading_Loss'].dropna()!=0])

    ### Fault distribution
    Total_faults = len(merged[merged['Fault']==merged['Fault']])
    Total_F0 = len(merged[merged['Fault']==0])
    Total_F1 = len(merged[merged['Fault']==1])
    Total_F2 = len(merged[merged['Fault']==2])
    Total_F3 = len(merged[merged['Fault']==3])
    Total_F4 = len(merged[merged['Fault']==4])
    Total_F5 = len(merged[merged['Fault']==5])

    F0 = np.round((Total_F0/Total_faults),3)*100
    F1 = np.round((Total_F1/Total_faults),3)*100
    F2 = np.round((Total_F2/Total_faults),3)*100
    F3 = np.round((Total_F3/Total_faults),3)*100
    F4 = np.round((Total_F4/Total_faults),3)*100
    F5 = np.round((Total_F5/Total_faults),3)*100
    SS = np.round((Fault_losses.mean()),3)*100
    SF = np.round(F2+F3+F4+F5,3)
    
    ### Fault distribution in time
    fault_distribution_fig = Plot_fault_distribution(merged)
    
    ### Save figures
    #folder = 'Figures/'
    #path = folder + str(meta['ID'].values[0]) + '/'
    #if not os.path.exists(path):
    #        os.mkdir(path)
    #f = path + str(meta['ID'].values[0])+'_inverter'+str(meta['Inverter'].values[0])+'_string'\
    #            +str(meta['String'].values[0]) +'_year'+str(meta['Year'].values[0])
    #f_cs = f + '_clear_sky_simulation.png'
    #f_fault = f + '_fault_heatmap.png'
    #f_current = f + '_current_loss_heatmap.png'
    #f_power = f + '_power_loss_heatmap.png'
    #f_dist = f + '_fault_distribution.png'
    
    #clear_sky_fig.write_image(f_cs, width=1000, height=500, scale=5)
    #fault_heatmap_fig.write_image(f_fault, width=1000, height=500, scale=5)
    #current_heatmap_fig.write_image(f_current, width=1000, height=500, scale=5)
    #power_heatmap_fig.write_image(f_power, width=1000, height=500, scale=5)
    #fault_distribution_fig.write_image(f_dist, width=1000, height=500, scale=5)
    
    ### Create word report
    #Create_word_report(meta, f_cs, f_fault, f_current, f_power, f_dist, F0, F1, F2, F3, F4, F5, SS, SF)
    
    return clear_sky_fig, fault_heatmap_fig, current_heatmap_fig, power_heatmap_fig,voltage_compare_fig, fault_distribution_fig, F0, F1, F2, F3, F4, F5, SS, SF

######################

def Create_word_report(meta, f_cs, f_fault, f_current, f_power, f_dist,F0, F1, F2, F3, F4, F5, SS, SF):
    
    from docx import Document
    from docx.shared import Pt
    from docx.shared import Inches

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    ### Add header and footer
    header = document.sections[0].header
    footer = document.sections[0].footer

    header_paragraph = header.paragraphs[0]
    logo_run = header_paragraph.add_run()
    logo_run.add_picture("Figures/00_Misc/logo_3S.png", width=Inches(1))

    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Created by Hugo Quest (hugo.quest@hotmail.com)"

    ### Title 
    document.add_heading('FDDA Report - Shading analysis results', 0)

    ### System meta data 
    document.add_heading('System metadata \n', level=1)

    table = document.add_table(rows=1, cols=10)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('ID').bold = True
    hdr_cells[1].paragraphs[0].add_run('Name').bold = True
    hdr_cells[2].paragraphs[0].add_run('Inverter').bold = True
    hdr_cells[3].paragraphs[0].add_run('String').bold = True
    hdr_cells[4].paragraphs[0].add_run('Capacity [kWp]').bold = True
    hdr_cells[5].paragraphs[0].add_run('Location').bold = True
    hdr_cells[6].paragraphs[0].add_run('Orientation [°]').bold = True
    hdr_cells[7].paragraphs[0].add_run('Tilt [°]').bold = True
    hdr_cells[8].paragraphs[0].add_run('Build date').bold = True
    hdr_cells[9].paragraphs[0].add_run('Studied year').bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = str(meta['ID'].values[0])
    row_cells[1].text = str(meta['Name'].values[0])
    row_cells[2].text = str(meta['Inverter'].values[0])
    row_cells[3].text = str(meta['String'].values[0])
    row_cells[4].text = str(meta['kWp'].values[0])
    row_cells[5].text = str(meta['Location'].values[0])
    row_cells[6].text = str(meta['Orientation'].values[0])
    row_cells[7].text = str(meta['Tilt'].values[0])
    row_cells[8].text = str(meta['Build_date'].values[0])
    row_cells[9].text = str(meta['Year'].values[0])
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)

    ### Part 1 - Shading quantification
    document.add_heading('A - Shading and Fault quantification', level=1)
    
    document.add_heading('1 | Table of fault classification \n', level=2)
    document.add_paragraph('Classification of faults defined in the FDDA. ' +
                            'The fault detection is based on thresholds around the DC voltage and current outputs, ' + \
                            'and the diagnosis depends on the combination of faults detected. The table below summarises ' +\
                            'the fault types differentiated by the algorithm, along with potential causes of the faults.' +\
                            ' F* faults are currently not identified explicitly.')
    # Table of fault classification
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Fault symbol').bold = True
    hdr_cells[1].paragraphs[0].add_run('Fault type').bold = True
    hdr_cells[2].paragraphs[0].add_run('Potential causes').bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = str('F0')
    row_cells[1].text = str('Normal')
    row_cells[2].text = str('')
    row_cells = table.add_row().cells
    row_cells[0].text = str('F1')
    row_cells[1].text = str('Cloudy')
    row_cells[2].text = str('Weather influence')
    row_cells = table.add_row().cells
    row_cells[0].text = str('F2')
    row_cells[1].text = str('Low V, Low I')
    row_cells[2].text = str('Connection fault, MPPT fault, Partial shading')
    row_cells = table.add_row().cells
    row_cells[0].text = str('F3')
    row_cells[1].text = str('Low V')
    row_cells[2].text = str('Shading with BPD activation, Module fault (short-circuited BPD)')
    row_cells = table.add_row().cells
    row_cells[0].text = str('F4')
    row_cells[1].text = str('High V, Low I')
    row_cells[2].text = str('Temperature effect due to shading, MPPT fault, Inverter shutdown')
    row_cells = table.add_row().cells
    row_cells[0].text = str('F5')
    row_cells[1].text = str('High V')
    row_cells[2].text = str('Temperature effect due to shading, MPPT fault, Inverter shutdown')
    row_cells = table.add_row().cells
    row_cells[0].text = str('F*')
    row_cells[1].text = str('Other')
    row_cells[2].text = str('Open-circuit fault, Soiling, Hot-spot, Potential-induced fault, Earth fault, Soldering, Discoloration, Cracks, Delamination,…')
    
    document.add_heading('2 | Fault analysis results \n', level=2)
    document.add_paragraph('Results of the FDDA for the studied system - shading factor, average power loss and \
                            relative occurences of faults.')
    p = document.add_paragraph()
    run = p.add_run('Definitions')
    run.bold = True
    document.add_paragraph('SF - Shading Factor [%]: percentage of power-producing time where shaded (F2-F5)', style='ListBullet')
    document.add_paragraph('SS - Shading Severity [%]: Average relative power loss during shading faults', style='ListBullet')
    document.add_paragraph('F0 - Relative occurence of F0 (normal) [%]', style='ListBullet')
    document.add_paragraph('F1 - Relative occurence of F1 (cloudy) [%]', style='ListBullet')
    document.add_paragraph('F2 - Relative occurence of F2 (low voltage, low current) [%]', style='ListBullet')
    document.add_paragraph('F3 - Relative occurence of F3 (low voltage) [%]', style='ListBullet')
    document.add_paragraph('F4 - Relative occurence of F4 (high voltage) [%]', style='ListBullet')
    document.add_paragraph('F5 - Relative occurence of F5 (high voltage, low current) [%]', style='ListBullet')

    # Table of fault types
    table = document.add_table(rows=1, cols=8)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('SF').bold = True
    hdr_cells[1].paragraphs[0].add_run('SS').bold = True
    hdr_cells[2].paragraphs[0].add_run('F0').bold = True
    hdr_cells[3].paragraphs[0].add_run('F1').bold = True
    hdr_cells[4].paragraphs[0].add_run('F2').bold = True
    hdr_cells[5].paragraphs[0].add_run('F3').bold = True
    hdr_cells[6].paragraphs[0].add_run('F4').bold = True
    hdr_cells[7].paragraphs[0].add_run('F5').bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = str(SF)
    row_cells[1].text = str(SS)
    row_cells[2].text = str(F0)
    row_cells[3].text = str(F5)
    row_cells[4].text = str(F1)
    row_cells[5].text = str(F2)
    row_cells[6].text = str(F3)
    row_cells[7].text = str(F4)
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(10)
    
    df_faults = pd.DataFrame(columns=['Fault Type', 'Occurence'])
    df_faults['Fault Type'] = ['F2', 'F3', 'F4', 'F5']
    df_faults['Occurence'] = [F1, F2, F3, F4]
    df_faults = df_faults.sort_values(by='Occurence', ascending=False).reset_index(drop=True)
    max_fault = df_faults['Fault Type'][0]
    max_fault_val = df_faults['Occurence'][0]
    document.add_paragraph('\nThe most prevalent fault in the system is ' + max_fault + ', which accounts for '\
                            + str(max_fault_val) + '% of the time where the system is at fault (excluding cloudy weather).')
    document.add_paragraph('When the system is at fault (F2-F5), the average relative power loss observed (compared to \
                            optimal clear-sky performance) is ' + str(SS) + '%.')

    ### Part 2 - Shading figures
    document.add_heading('B - FDDA Figure outputs', level=1)

    document.add_heading('1 | Clear-sky power simulation (single-diode model) \n', level=2)
    document.add_paragraph('Clear-sky DC power output simulation based on single-diode model compared to \
                            the actual DC power output.')
    document.add_picture(f_cs, width=Inches(6), height=Inches(3))
    
    document.add_page_break()
    
    document.add_heading('2 | Shading fault and Current loss heatmap \n', level=2)
    document.add_paragraph('Shading faults, relative current loss and relative power loss heatmaps. These figures show the ' +\
                            'temporal distribution of faults and losses in time (x-axis corresponds to hours in a day, ' + \
                            'y-axis corresponds to the date). The relative losses are computed by comparing with the clear-sky ' +\
                            'optimal outputs.')
    document.add_picture(f_fault, width=Inches(6), height=Inches(3))
    document.add_picture(f_current, width=Inches(6), height=Inches(3))
    document.add_picture(f_power, width=Inches(6), height=Inches(3))
        
    document.add_heading('3 | Shading fault occurence histogram \n', level=2)
    document.add_paragraph('Histogram of fault occurence as a function of the relative time of day. The values of the x-axis represent ' + \
                           'normalised hours in a day, given the differences in day length during the year.')
    document.add_picture(f_dist, width=Inches(6), height=Inches(3))

    ### Save document
    folder = 'Outputs/'
    path = folder + str(meta['ID'].values[0]) + '/'
    if not os.path.exists(path):
            os.mkdir(path)
    document.save(path + 'FDDA_output_' + str(meta['ID'].values[0]) + '_inv' + str(meta['Inverter'].values[0]) 
                       + '_string' + str(meta['String'].values[0]) + '_year' 
                       + str(meta['Year'].values[0]) + '.docx')
    
    
  
